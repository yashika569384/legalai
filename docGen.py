import os
import re
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from groq import Groq

load_dotenv()

client = Groq(api_key=os.getenv("GROQ_API_KEY"))
GROQ_MODEL = "llama-3.3-70b-versatile"

DRAFTING_SYSTEM_PROMPT = """
You are an expert Indian legal document drafter with deep knowledge of Indian law,
the Indian Penal Code, Civil Procedure Code, Transfer of Property Act, Consumer Protection Act,
and all other relevant Indian legislation.

Your task is to draft a complete, professional, and legally sound document based on the user's request.

Guidelines:
- Draft a COMPLETE document.
- Use clear placeholders for user-editable fields like [NAME], [DATE], [ADDRESS], [AMOUNT], [SIGNATURE].
- Use placeholders wherever personal or case-specific information is required so the user can easily edit the final document.
- Use formal legal language appropriate for Indian courts and legal practice.
- Include all standard clauses, sections, and legal language relevant to the document type.
- Structure the document properly with headings, numbered clauses, and signature blocks.
- Reference relevant Indian laws, sections, and acts where appropriate.
- The document should be ready to use with only minor edits needed for specific details.
- Do NOT add any commentary, explanation, or preamble — output ONLY the document text itself.
"""


def _add_paragraph(doc, text, bold=False, center=False, font_size=12):
    """Add a paragraph with plain black text — no blue heading styles."""
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)  # always black
    return p


def _is_heading(line: str) -> bool:
    """Detect lines that should be bold headings."""
    stripped = line.strip()
    # All caps line (e.g. LEGAL NOTICE, IN THE COURT OF...)
    if stripped.isupper() and len(stripped) > 3:
        return True
    # Short line ending with colon (e.g. "Subject:", "From:", "To:")
    if stripped.endswith(":") and len(stripped) < 60:
        return True
    return False


def _is_center_heading(line: str) -> bool:
    """Lines that should be centered (court name, document title)."""
    stripped = line.strip()
    center_keywords = [
        "in the court", "hon'ble", "legal notice", "bail application",
        "power of attorney", "lease agreement", "affidavit", "employment contract",
        "cease and desist", "sale deed", "partnership deed"
    ]
    return any(kw in stripped.lower() for kw in center_keywords)


def generate_legal_document(prompt: str, save_dir: str = 'static/generated_docs'):
    """
    Uses Groq AI to draft any legal document the user describes.
    Fully open-ended — no restrictions on document type.
    """
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)

    # Step 1: Draft the document
    response = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[
            {"role": "system", "content": DRAFTING_SYSTEM_PROMPT.strip()},
            {"role": "user", "content": f"User Request: {prompt.strip()}\n\nDraft the complete legal document now:"}
        ],
        temperature=0.3,
        max_tokens=2048
    )
    document_text = response.choices[0].message.content.strip()

    if not document_text:
        raise ValueError("Could not generate the document. Please try rephrasing your request.")

    # Step 2: Generate filename
    file_name = _generate_filename(prompt)

    # Step 3: Build Word document with proper plain black formatting
    doc = DocxDocument()

    # Set default font for the document
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)

    for line in document_text.split("\n"):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        if _is_center_heading(stripped):
            _add_paragraph(doc, stripped, bold=True, center=True, font_size=12)
        elif _is_heading(stripped):
            _add_paragraph(doc, stripped, bold=True, center=False, font_size=12)
        else:
            _add_paragraph(doc, stripped, bold=False, center=False, font_size=12)

    file_path = os.path.join(save_dir, file_name)
    doc.save(file_path)
    return file_path, file_name


def _generate_filename(prompt: str) -> str:
    """Ask Groq for a snake_case filename, fall back to prompt words."""
    try:
        resp = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": (
                        "Based on the following user request for a legal document, "
                        "reply with ONLY a short snake_case filename (lowercase, underscores, no extension, max 5 words). "
                        "Examples: legal_notice_landlord, bail_application, nda_employee, divorce_petition.\n\n"
                        f"User request: {prompt.strip()}"
                    )
                }
            ],
            temperature=0,
            max_tokens=20
        )
        raw = resp.choices[0].message.content.strip().lower()
        clean = re.sub(r'[^a-z0-9_]', '_', raw).strip('_')
        clean = re.sub(r'_+', '_', clean)
        if clean:
            return f"{clean}.docx"
    except Exception:
        pass

    # Fallback
    words = re.sub(r'[^a-z0-9\s]', '', prompt.lower()).split()
    stopwords = {
        'a', 'an', 'the', 'for', 'of', 'to', 'in', 'on', 'and', 'or',
        'write', 'draft', 'create', 'generate', 'make', 'me', 'my', 'i',
        'please', 'need', 'want', 'legal', 'document'
    }
    meaningful = [w for w in words if w not in stopwords][:4]
    return ('_'.join(meaningful) or 'legal_document') + '.docx'